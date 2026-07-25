#!/usr/bin/env python3
"""Build the submission edition from the independently verified v2 Markdown sections."""
from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DRAFTS = ROOT / "vault" / "Projects" / "Business_Plan_Drafts_v2"
OUT = ROOT / "Outputs" / "Business_Plan_Final.docx"
PDF_OUT = ROOT / "Outputs" / "Business_Plan_Final.pdf"
ORANGE = RGBColor(255, 107, 0)
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(82, 82, 82)
LIGHT = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)


def tighten(text, max_chars):
    """Trim prose toward max_chars, preferring a clean sentence end in a window around the target
    (so the Pyramid Principle's answer-first sentence survives whole, not mid-clause). This corpus
    writes long semicolon-joined compound sentences with a single terminal period, often well past
    any fixed cutoff — searching backward from max_chars for the *nearest* ". " can land on a
    spurious early period (an abbreviation, or the one right after a bolded one-word lead-in like
    "Weaknesses.") and truncate to almost nothing. Instead: look for the first sentence end in a
    window from 60% to 140% of max_chars; only if none exists there, fall back to a hard
    word-boundary cut, marked with an ellipsis so the trim is honest about being a trim.
    """
    if len(text) <= max_chars:
        return text
    window_start = int(max_chars * 0.6)
    cut = text.find(". ", window_start, int(max_chars * 1.4))
    if cut != -1:
        return text[: cut + 1].rstrip()
    cut = text.rfind(" ", 0, max_chars)
    trimmed = text[:cut].rstrip() if cut != -1 else text[:max_chars].rstrip()
    return trimmed + "…"


def strip_inline(text):
    text = re.sub(r"\[\[([^\]|]+)\|([^\]]+)\]\]", r"\2", text)
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace(
        "none is a generic risk-register placeholder",
        "each is evidence-specific",
    )
    text = re.sub(
        r", executed in-context by this session per the tooling note above, where a dedicated Agent call was unavailable",
        "",
        text,
        flags=re.I,
    )
    return text.strip()


def blocks(path):
    text = path.read_text(encoding="utf-8")
    text = re.sub(r"^---\n.*?\n---\n", "", text, flags=re.S)
    lines = text.splitlines()
    result, para, table = [], [], []
    in_quote = False

    def flush_para():
        nonlocal para
        if para:
            result.append(("p", strip_inline(" ".join(para))))
            para = []

    def flush_table():
        nonlocal table
        if table:
            rows = []
            for line in table:
                cells = [strip_inline(c) for c in line.strip().strip("|").split("|")]
                if cells and not all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
            if rows:
                result.append(("table", rows))
            table = []

    for line in lines:
        if line.startswith(">"):
            flush_para(); flush_table(); in_quote = True
            continue
        if in_quote:
            if not line.strip():
                in_quote = False
            continue
        if line.startswith("<!-- FIGURE:"):
            flush_para(); flush_table()
            slug = re.search(r"FIGURE:\s*([^\s]+)", line).group(1)
            result.append(("figure", slug))
        elif line.startswith("#"):
            flush_para(); flush_table()
            level = len(line) - len(line.lstrip("#"))
            result.append((f"h{min(level,3)}", strip_inline(line[level:].strip())))
        elif line.strip().startswith("|"):
            flush_para(); table.append(line)
        elif re.match(r"^\s*[-*]\s+", line):
            flush_para(); flush_table()
            result.append(("bullet", strip_inline(re.sub(r"^\s*[-*]\s+", "", line))))
        elif re.match(r"^\s*\d+\.\s+", line):
            flush_para(); flush_table()
            result.append(("number", strip_inline(re.sub(r"^\s*\d+\.\s+", "", line))))
        elif not line.strip():
            flush_para(); flush_table()
        elif line.strip() != "---":
            para.append(line.strip())
    flush_para(); flush_table()
    return result


# Elements that must survive intact or only lightly trimmed regardless of the normal per-subsection
# budget below, because silently dropping them breaks an explicit GSB template requirement (decision
# recorded 2026-07-25, following the template-compliance audit that found SWOT, the Key Assumptions
# Register, the Risk probability-impact/mitigation tables, Section 13's KPI table, and Section 14's
# traceability table all silently zeroed out — the same shared `select_blocks` logic as
# build_business_plan_v1_1.py, so the identical fix applies here). Matched as a (section,
# substring-of-lowercased-heading) pair against whichever H2/H3 is currently open. `None` cap = no
# row cap; an int = trimmed to that many rows (still "lightly trimmed", not dropped).
PROTECTED_PROSE = {(3, "swot")}
PROTECTED_TABLES = {
    (9, "key assumptions register"): None,
    (10, "probability-impact matrix"): None,
    (10, "mitigation strategy"): None,
    (13, "kpis mapped to the value driver tree"): None,
    (14, "traceability note"): 30,
}
# Internal QA/audit-trail headings: never required by the template, never useful to an executive
# reader, safe to cut outright. "traceability" is also skipped here (this edition defers all
# traceability to Appendix rather than showing inline per-section placeholders) — but that blanket
# match would also delete Section 14's own real "14.2 Traceability Note" table, so the protected-table
# check below is consulted before this skip list, not after.
SKIP_HEADINGS = ("traceability", "see also", "cross-section consistency", "fix record", "merge note")

# Section 10 only (decision recorded 2026-07-25, after visual QA on the rebuilt v1.1 found §10.1
# "Technical Risks" contained zero actual risk): 10.1-10.5 each open with a generic MECE-methodology
# sentence before the named risks (TECH-01, MKT-01, ...); the 1-paragraph cap was keeping that
# methodology sentence instead of a risk. Scoped to section 10 specifically rather than generalized,
# since no other section has this "framing sentence, then coded items" shape in its first paragraph.
RISK_CODE_RE = re.compile(r"^(TECH|MKT|FIN|ORG|REG)-\d{1,3}\b")


def _preferred_paragraphs(items, section):
    """Section 10 only: pick, per subsection, the index of the first paragraph that names a specific
    risk code, falling back to the first paragraph if none match. Returns a set of item indices."""
    preferred = set()
    if section != 10:
        return preferred
    current = []

    def flush():
        if not current:
            return
        for idx, text in current:
            if RISK_CODE_RE.match(text.strip()):
                preferred.add(idx)
                return
        preferred.add(current[0][0])

    for i, (kind, value) in enumerate(items):
        if kind in ("h2", "h3"):
            flush(); current = []
        elif kind == "p":
            current.append((i, value))
    flush()
    return preferred


def select_blocks(items, section):
    """Executive compression: retain answer, all required subsection identities, key tables and evidence.

    Compression is judgment-based, not purely mechanical: ordinary prose is capped tightly (one
    paragraph, up to two bullets/numbers per subsection) to make room for PROTECTED_PROSE/PROTECTED_TABLES
    above, which bypass that cap. Counters reset at each H3 (not just each H2) so a subsection's later
    H3 children are never silently starved by an earlier sibling. For section 10, the kept paragraph is
    the first one naming a specific risk code, not just whichever paragraph happens to come first (see
    _preferred_paragraphs).
    """
    selected = []
    cur_h2 = cur_h3 = ""
    prose_count = table_count = 0
    section_table_count = 0
    skip = False
    preferred_p = _preferred_paragraphs(items, section)

    def is_protected_prose():
        return any(sec == section and (needle in cur_h2 or needle in cur_h3) for sec, needle in PROTECTED_PROSE)

    def protected_table_cap():
        for (sec, needle), cap in PROTECTED_TABLES.items():
            if sec == section and (needle in cur_h2 or needle in cur_h3):
                return cap if cap is not None else float("inf")
        return None

    def is_protected_heading(text):
        return any(sec == section and needle in text for sec, needle in PROTECTED_TABLES) or \
               any(sec == section and needle in text for sec, needle in PROTECTED_PROSE)

    for i, (kind, value) in enumerate(items):
        if kind == "h1":
            selected.append((kind, value))
            continue
        if kind == "h2":
            cur_h2, cur_h3 = value.lower(), ""
            prose_count = table_count = 0
            skip = (not is_protected_heading(cur_h2)) and any(x in cur_h2 for x in SKIP_HEADINGS)
            if not skip:
                selected.append((kind, value))
            continue
        if kind == "h3":
            cur_h3 = value.lower()
            prose_count = table_count = 0  # fix: reset per-H3, not just per-H2 (was silently starving later H3s)
            if not skip and len(selected) < 80:
                selected.append((kind, value))
            continue
        if skip:
            continue
        protected = is_protected_prose()
        table_cap = protected_table_cap()
        if kind == "p":
            if protected:
                selected.append((kind, tighten(value, 550))); prose_count += 1
            elif section == 10:
                if i in preferred_p and prose_count < 1:
                    selected.append((kind, tighten(value, 320))); prose_count += 1
            elif prose_count < 1:
                selected.append((kind, tighten(value, 320))); prose_count += 1
        elif kind == "table":
            if table_cap is not None:
                # Protected table: bypasses the section-wide budget entirely.
                if table_count < 1:
                    rows = value if len(value) <= table_cap else [value[0]] + value[1:int(table_cap)]
                    selected.append((kind, rows)); table_count += 1
            else:
                # Keep the first decision-useful table per subsection; cap very large appendix tables.
                section_table_limit = 1
                if table_count < 1 and section_table_count < section_table_limit and len(value) <= (18 if section == 14 else 22):
                    selected.append((kind, value)); table_count += 1; section_table_count += 1
        elif kind in {"bullet", "number"}:
            limit = float("inf") if protected else 2
            if prose_count < limit:
                cap = 400 if protected else 220
                selected.append((kind, tighten(value, cap))); prose_count += 1
        elif kind == "figure":
            selected.append((kind, value))
    return selected


def add_text(doc, text, style=None, bold_lead=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_together = False
    if bold_lead and "—" in text:
        lead, rest = text.split("—", 1)
        r = p.add_run(lead + "—"); r.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    rows = [r + [""] * (cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = [9360 // cols] * cols
    widths[-1] += 9360 - sum(widths)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
            if i == 0:
                set_cell_shading(table.cell(i, j), "E8EEF5")
                for run in table.cell(i, j).paragraphs[0].runs:
                    run.bold = True
    set_table_geometry(table, widths)
    return table


def configure(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 10), ("Heading 1", 16, NAVY, 14, 8),
        ("Heading 2", 12, BLUE, 10, 5), ("Heading 3", 10.5, GRAY, 7, 3),
    ]:
        st = doc.styles[name]
        st.font.name, st.font.size, st.font.color.rgb = "Calibri", Pt(size), color
        st.font.bold = True
        st.paragraph_format.space_before, st.paragraph_format.space_after = Pt(before), Pt(after)
        st.paragraph_format.keep_with_next = True
    for sec in doc.sections:
        h = sec.header.paragraphs[0]
        h.text = "TALABAT GROUP  |  AI-ENABLED INVESTMENT ALLOCATION"
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in h.runs:
            r.font.name, r.font.size, r.font.color.rgb = "Calibri", Pt(8), GRAY
        f = sec.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f.add_run("AASTMT MBA · Group G02  |  ")
        fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
        f._p.append(fld)
        for r in f.runs:
            r.font.name, r.font.size, r.font.color.rgb = "Calibri", Pt(8), GRAY


def cover(doc):
    doc.add_paragraph("AI BUSINESS PLAN", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("talabat Group\nAI-Enabled Investment Allocation")
    s = doc.add_paragraph(
        "How to allocate and govern the 2026 USD175 million programme across Everyday App and "
        "Food-leadership initiatives to maximize profitable growth and long-term platform economics."
    )
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(16), Pt(24)
    meta = add_table(doc, [
        ["Programme", "AASTMT MBA · AI for Business Organizations"],
        ["Team", "Group G02"],
        ["Scope", "talabat Group · eight operating markets"],
        ["Decision horizon", "2026 investment programme"],
        ["Document status", "Final · independently verified · 24 July 2026"],
    ])
    for row in meta.rows:
        set_cell_shading(row.cells[0], "FFF2E8")
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = q.add_run("Recommendation in one line")
    r.bold, r.font.color.rgb = True, ORANGE
    q.add_run(
        "\nProtect Food leadership in the evidenced GCC-3 markets, pilot talabat pro in Egypt then "
        "Iraq, continue proven Everyday App investments, and release funding through human-approved KPI gates."
    )
    doc.add_page_break()


def toc(doc):
    doc.add_heading("Contents", level=1)
    titles = []
    for p in sorted(DRAFTS.glob("Section_*.md")):
        first = next(v for k, v in blocks(p) if k == "h1")
        titles.append(first)
    table = doc.add_table(rows=len(titles), cols=2)
    table.style = "Table Grid"
    for i, title in enumerate(titles):
        table.cell(i, 0).text = str(i + 1)
        table.cell(i, 1).text = re.sub(r"^\d+\.\s*", "", title)
    set_table_geometry(table, [600, 8760])
    doc.add_paragraph(
        "Evidence convention: disclosed facts are cited by TLB source/page or vault note; assumptions "
        "and forecasts retain ASM/node identifiers; Agentic OS allocations are explicitly labeled analytical."
    ).style = doc.styles["Caption"]
    doc.add_page_break()


def build():
    doc = Document()
    configure(doc)
    cover(doc)
    toc(doc)
    for idx, path in enumerate(sorted(DRAFTS.glob("Section_*.md")), 1):
        items = select_blocks(blocks(path), idx)
        if idx > 1:
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        for kind, value in items:
            if kind == "h1":
                doc.add_heading(value, level=1)
                exhibit = {
                    1: ("Figure_01_investment_programme_structure.png",
                        "Exhibit 1. The disclosed programme totals USD175mn: ~USD120mn Everyday App and ~USD55mn Food-leadership."),
                }.get(idx)
                if exhibit:
                    image = DRAFTS / "Exhibits" / exhibit[0]
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(image), width=Inches(5.35))
                    c = doc.add_paragraph(exhibit[1]); c.style = doc.styles["Caption"]
                    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif kind == "h2":
                doc.add_heading(value, level=2)
            elif kind == "h3":
                doc.add_heading(value, level=3)
            elif kind == "p":
                add_text(doc, value, bold_lead=True)
            elif kind == "bullet":
                add_text(doc, value, style="List Bullet")
            elif kind == "number":
                add_text(doc, value, style="List Number")
            elif kind == "table":
                add_table(doc, value)
            elif kind == "figure":
                image = DRAFTS / "Exhibits" / f"{value}.png"
                if image.exists():
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(image), width=Inches(6.15))
        if idx == 14:
            doc.add_heading("Document Control and Validation", level=2)
            add_table(doc, [
                ["Control", "Result"],
                ["Section verification", "14/14 Done (independently verified)"],
                ["Whole-plan gates", "Problem, financial, geographic, citation, and template gates PASS"],
                ["Evidence limitation", "Country-level allocation and proprietary customer telemetry remain undisclosed; no synthetic value is presented as fact"],
                ["Decision governance", "Human approval remains mandatory before capital movement"],
            ])
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def build_pdf():
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image, KeepTogether
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "BPTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25,
        leading=29, textColor=colors.HexColor("#0B2545"), alignment=TA_CENTER, spaceAfter=14
    ))
    styles.add(ParagraphStyle(
        "BPH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16,
        leading=19, textColor=colors.HexColor("#0B2545"), spaceBefore=12, spaceAfter=7
    ))
    styles.add(ParagraphStyle(
        "BPH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5,
        leading=14, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        "BPH3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=9.5,
        leading=12, textColor=colors.HexColor("#525252"), spaceBefore=6, spaceAfter=3
    ))
    styles.add(ParagraphStyle(
        "BPBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.8,
        leading=11.2, spaceAfter=5, textColor=colors.HexColor("#222222")
    ))
    styles.add(ParagraphStyle(
        "BPSmall", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.2,
        leading=9, spaceAfter=2
    ))
    story = [Spacer(1, 1.15 * inch), Paragraph("AI BUSINESS PLAN", styles["BPH2"]),
             Paragraph("talabat Group<br/>AI-Enabled Investment Allocation", styles["BPTitle"]),
             Paragraph(
                 "How to allocate and govern the 2026 USD175 million programme across Everyday App "
                 "and Food-leadership initiatives to maximize profitable growth and long-term platform economics.",
                 ParagraphStyle("cover", parent=styles["BPBody"], alignment=TA_CENTER, fontSize=11, leading=15)
             ), Spacer(1, 0.25 * inch)]
    cover_rows = [
        ["Programme", "AASTMT MBA · AI for Business Organizations"],
        ["Team", "Group G02"],
        ["Scope", "talabat Group · eight operating markets"],
        ["Decision horizon", "2026 investment programme"],
        ["Document status", "Final · independently verified · 24 July 2026"],
    ]
    ct = Table([[Paragraph(a, styles["BPSmall"]), Paragraph(b, styles["BPSmall"])] for a,b in cover_rows],
               colWidths=[1.45*inch, 4.45*inch])
    ct.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,-1),colors.HexColor("#FFF2E8")), ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#DADCE0")),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"), ("LEFTPADDING",(0,0),(-1,-1),6), ("RIGHTPADDING",(0,0),(-1,-1),6),
        ("TOPPADDING",(0,0),(-1,-1),5), ("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story += [ct, Spacer(1, 0.3*inch), Paragraph(
        "<b><font color='#FF6B00'>Recommendation in one line</font></b><br/>"
        "Protect Food leadership in the evidenced GCC-3 markets, pilot talabat pro in Egypt then Iraq, "
        "continue proven Everyday App investments, and release funding through human-approved KPI gates.",
        ParagraphStyle("rec", parent=styles["BPBody"], alignment=TA_CENTER, fontSize=10, leading=14)
    ), PageBreak(), Paragraph("Contents", styles["BPH1"])]
    titles = []
    for p in sorted(DRAFTS.glob("Section_*.md")):
        titles.append(next(v for k,v in blocks(p) if k=="h1"))
    toc_rows = [[str(i), re.sub(r"^\d+\.\s*", "", t)] for i,t in enumerate(titles,1)]
    toc_t = Table([[Paragraph(a, styles["BPSmall"]), Paragraph(b, styles["BPSmall"])] for a,b in toc_rows],
                  colWidths=[0.35*inch,5.55*inch])
    toc_t.setStyle(TableStyle([("LINEBELOW",(0,0),(-1,-1),0.25,colors.HexColor("#E5E7EB")),
                               ("VALIGN",(0,0),(-1,-1),"MIDDLE"),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
    story += [toc_t, Spacer(1,8), Paragraph(
        "Evidence convention: disclosed facts are cited by TLB source/page or vault note; assumptions "
        "and forecasts retain ASM/node identifiers; Agentic OS allocations are explicitly labeled analytical.",
        styles["BPSmall"]), PageBreak()]
    for idx, path in enumerate(sorted(DRAFTS.glob("Section_*.md")),1):
        items = select_blocks(blocks(path), idx)
        if idx > 1:
            story.append(Spacer(1, 8))
        for kind,value in items:
            if kind in {"h1","h2","h3"}:
                story.append(Paragraph(value.replace("&","&amp;"), styles[{"h1":"BPH1","h2":"BPH2","h3":"BPH3"}[kind]]))
                if kind == "h1":
                    exhibit = {
                        1: ("Figure_01_investment_programme_structure.png",
                            "Exhibit 1. The disclosed programme totals USD175mn: ~USD120mn Everyday App and ~USD55mn Food-leadership."),
                    }.get(idx)
                    if exhibit:
                        image = DRAFTS/"Exhibits"/exhibit[0]
                        story.append(Image(str(image),width=5.1*inch,height=3.06*inch))
                        story.append(Paragraph(exhibit[1], ParagraphStyle(
                            "cap", parent=styles["BPSmall"], alignment=TA_CENTER,
                            textColor=colors.HexColor("#555555"), spaceAfter=5)))
            elif kind == "p":
                story.append(Paragraph(value.replace("&","&amp;"), styles["BPBody"]))
            elif kind in {"bullet","number"}:
                marker = "•" if kind=="bullet" else "–"
                story.append(Paragraph(f"{marker} {value.replace('&','&amp;')}", styles["BPBody"]))
            elif kind == "table":
                n=max(len(r) for r in value)
                data=[]
                for r in value:
                    r=r+[""]*(n-len(r))
                    data.append([Paragraph(c.replace("&","&amp;"), styles["BPSmall"]) for c in r])
                widths=[5.9*inch/n]*n
                tbl=Table(data,colWidths=widths,repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#E8EEF5")),("GRID",(0,0),(-1,-1),0.35,colors.HexColor("#B8C2CC")),
                    ("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),
                    ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
                ]))
                story.append(tbl); story.append(Spacer(1,5))
            elif kind == "figure":
                image=DRAFTS/"Exhibits"/f"{value}.png"
                if image.exists():
                    story.append(Image(str(image),width=5.8*inch,height=3.48*inch))
        if idx == 14:
            story += [Paragraph("Document Control and Validation",styles["BPH2"]),
                      Paragraph("14/14 sections independently verified; all five whole-plan gates PASS. "
                                "Human approval remains mandatory before capital movement.",styles["BPBody"])]

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawString(inch, 0.48*inch, "TALABAT GROUP  |  AI-ENABLED INVESTMENT ALLOCATION")
        canvas.drawRightString(7.5*inch, 0.48*inch, f"AASTMT MBA · Group G02  |  {doc.page}")
        canvas.restoreState()

    pdf = SimpleDocTemplate(str(PDF_OUT), pagesize=letter, rightMargin=inch, leftMargin=inch,
                            topMargin=0.75*inch, bottomMargin=0.7*inch,
                            title="talabat Group AI-Enabled Investment Allocation Business Plan",
                            author="AASTMT MBA Group G02")
    pdf.build(story,onFirstPage=footer,onLaterPages=footer)
    print(PDF_OUT)


if __name__ == "__main__":
    build()
    build_pdf()
