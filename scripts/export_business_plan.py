#!/usr/bin/env python3
"""Assemble the talabat Group Business Plan's Section_XX_*.md drafts into a submission-ready DOCX and
PDF. Presentation only — never alters a number, claim, citation, or conclusion found in the source
Markdown (mirrors executive-document-formatting/SKILL.md's content-protection rules).

DOCX is built natively via python-docx (real paragraph/heading styles, a native Word TOC field, tables,
embedded figures). PDF is produced by rendering the same parsed content as HTML and printing it with
headless Chromium (Playwright) — **not** LibreOffice/soffice: this environment's LibreOffice
installation was tested directly during the 2026-07-23 Final Execution Readiness pass and fails to load
*any* document, including a single-paragraph trivial one ("Error: source file could not be loaded",
reproduced with a fresh profile and outside the sandbox) — the same disclosed failure mode recorded for
the original 22/07/2026 export, now re-confirmed rather than assumed fixed. Headless Chromium is
confirmed working in this environment (this repo's standard pre-installed browser, per the session's
own environment notes) and is the same substitute method the original export used, kept consistent
rather than reintroducing a two-pipeline DOCX/PDF divergence.

Two modes:
  --gate (default)   Refuses to run unless all 14 Section_XX_*.md drafts exist and carry
                      `status: ... Done (independently verified)` in frontmatter. Writes to Outputs/.
  --smoke-test        Bypasses the gate to verify the pipeline mechanically works. Stamps every page
                      with a "CAPABILITY VERIFICATION ONLY" watermark and writes to
                      vault/Validation/_export_smoke_test/ instead of Outputs/ — never produces
                      anything that could be mistaken for a real submission.

Usage:
  python3 scripts/export_business_plan.py --smoke-test
  python3 scripts/export_business_plan.py            # real export, once all 14 sections are Done

Requires: python-docx, playwright (pip) with a Chromium executable — this environment's pre-installed
one at /opt/pw-browsers is auto-detected; see find_chromium() below.
"""
import argparse
import glob
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DRAFTS_DIR = os.path.join(REPO_ROOT, "vault", "Projects", "Business_Plan_Drafts_v2")
EXHIBITS_DIR = os.path.join(DRAFTS_DIR, "Exhibits")
MANIFEST_PATH = os.path.join(EXHIBITS_DIR, "Exhibits_Manifest.md")
TRACKER_PATH = os.path.join(
    REPO_ROOT, "vault", "Projects", "Talabat-Group-AI-Investment-Allocation-Business-Plan.md"
)
OUTPUTS_DIR = os.path.join(REPO_ROOT, "Outputs")
SMOKE_TEST_DIR = os.path.join(REPO_ROOT, "vault", "Validation", "_export_smoke_test")
GENERATION_CONTRACT_PATH = os.path.join(REPO_ROOT, "Business_Plan_Generation_Contract.md")
EDITORIAL_STANDARD_PATH = os.path.join(
    REPO_ROOT, "vault", "Architecture", "Business_Plan_Editorial_Standard.md"
)
PREPUBLICATION_GATE_PATHS = [
    ("Factual QA", os.path.join(REPO_ROOT, "vault", "Validation", "Citation_Audit_Whole_Plan_v2_Pass4.md")),
    ("Financial QA", os.path.join(REPO_ROOT, "vault", "Validation", "Financial_Integrity_Gate.md")),
    ("Geographic QA", os.path.join(REPO_ROOT, "vault", "Validation", "Geographic_Evidence_Gate.md")),
    ("Decision consistency QA", os.path.join(REPO_ROOT, "vault", "Validation", "Decision_Consistency_QA.md")),
    ("Editorial readability QA", os.path.join(REPO_ROOT, "vault", "Validation", "Editorial_Readability_QA.md")),
    ("Content completeness QA", os.path.join(REPO_ROOT, "vault", "Validation", "Content_Completeness_QA.md")),
    ("Template compliance QA", os.path.join(REPO_ROOT, "vault", "Validation", "Template_Compliance_Checklist.md")),
    ("External-reader test", os.path.join(REPO_ROOT, "vault", "Validation", "External_Reader_Test.md")),
]

ORANGE = RGBColor(0xFF, 0x6B, 0x00)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)

FIGURE_MARKER_RE = re.compile(r"<!--\s*FIGURE:\s*(Figure_\d+_[A-Za-z0-9_]+)\s*-->")
RAW_INTERNAL_REFERENCE_RE = re.compile(
    r"(?:\bvault[/\\]|\b[\w.-]+\.md\b|\b(?:DEC|ASM)-\d{3}\b|"
    r"\b(?:bp-orchestrator|qa-review-agent|research-agent|forecasting-agent|"
    r"decision-steward|evidence-citation-agent|kpi-agent|exec-summary-agent|"
    r"business-plan-drafting|executive-document-formatting)\b)",
    re.IGNORECASE,
)
PROHIBITED_EXPRESSIONS = (
    "corpus", "ghost deck", "middle path", "not re-litigated", "not one it invents",
    "stated honestly", "not merely asserted", "does not fabricate", "does not paper over",
    "deliberate framing choice, not an oversight", "comparative in kind, not degree",
    "this os's own proposal", "this section's own contribution", "load-bearing claim",
    "evidentiary anchor",
)


# ---------------------------------------------------------------------------
# Markdown -> intermediate blocks
# ---------------------------------------------------------------------------

def strip_frontmatter(text):
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            fm = text[3:end].strip()
            body = text[end + 4:]
            meta = {}
            for line in fm.splitlines():
                if ":" in line:
                    k, _, v = line.partition(":")
                    meta[k.strip()] = v.strip().strip('"')
            return meta, body
    return {}, text


def editorial_preflight(section_files):
    """Return blocking Version 1.2 contract/editorial findings before a real export."""
    failures = []
    for label, path in (
        ("Business Plan Generation Contract", GENERATION_CONTRACT_PATH),
        ("Business Plan Editorial Standard", EDITORIAL_STANDARD_PATH),
    ):
        if not os.path.exists(path):
            failures.append(f"{label}: mandatory authority file is missing")

    for label, path in PREPUBLICATION_GATE_PATHS:
        if not os.path.exists(path):
            failures.append(f"{label}: required validation artifact is missing")
            continue
        report = open(path, encoding="utf-8").read()
        if not re.search(r"\bPASS\b", report, re.IGNORECASE):
            failures.append(f"{label}: validation artifact does not record PASS")

    section_one_body = ""
    for path in section_files:
        section_number = int(re.search(r"Section_(\d+)", os.path.basename(path)).group(1))
        _, body = strip_frontmatter(open(path, encoding="utf-8").read())
        if section_number == 1:
            section_one_body = body
        if section_number > 13:  # Technical/internal traceability is allowed in labelled appendices.
            continue

        for line_no, line in enumerate(body.splitlines(), 1):
            if RAW_INTERNAL_REFERENCE_RE.search(line):
                failures.append(
                    f"{os.path.basename(path)}:{line_no}: raw internal reference in the main body"
                )
            lowered = line.lower().replace("’", "'")
            for phrase in PROHIBITED_EXPRESSIONS:
                if phrase in lowered:
                    failures.append(
                        f"{os.path.basename(path)}:{line_no}: prohibited expression '{phrase}'"
                    )
            if re.match(r"^\s*[-*]\s+", line) and re.search(
                r"\b(?:and|or|of|to|for|from)\s*[.;:]?\s*$", line, re.IGNORECASE
            ):
                failures.append(
                    f"{os.path.basename(path)}:{line_no}: bullet appears to end mid-thought"
                )

        prose = re.sub(r"```.*?```", "", body, flags=re.DOTALL)
        for sentence in re.split(r"(?<=[.!?])\s+", prose):
            clean = re.sub(r"[#>*_`|[\]()-]", " ", sentence)
            words = re.findall(r"\b[\w’'-]+\b", clean)
            if len(words) > 35:
                excerpt = " ".join(words[:10])
                failures.append(
                    f"{os.path.basename(path)}: sentence over 35 words ('{excerpt}…')"
                )

    if "Evidence Basis and Limitations" not in section_one_body:
        failures.append(
            "Section 1: required consolidated 'Evidence Basis and Limitations' section is missing"
        )
    return failures


def parse_blocks(body):
    """Very small Markdown block parser: headings, blockquotes, tables, bullet lists, figure
    markers, horizontal rules, and plain paragraphs. Good enough for this repo's own drafting style
    (checked against Section_02/09's actual structure), not a general CommonMark implementation."""
    lines = body.splitlines()
    blocks = []
    i = 0
    para_buf = []

    def flush_para():
        if para_buf:
            blocks.append(("para", " ".join(para_buf).strip()))
            para_buf.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            flush_para()
            i += 1
            continue

        m = FIGURE_MARKER_RE.search(stripped)
        if m:
            flush_para()
            blocks.append(("figure", m.group(1)))
            i += 1
            continue

        if stripped == "---":
            flush_para()
            blocks.append(("hr", None))
            i += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading_match:
            flush_para()
            level = len(heading_match.group(1))
            blocks.append(("heading", (level, heading_match.group(2).strip())))
            i += 1
            continue

        if stripped.startswith(">"):
            flush_para()
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append(("quote", " ".join(quote_lines)))
            continue

        if stripped.startswith("|"):
            flush_para()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\|?[\s:|-]+\|?$", row)
            ]
            blocks.append(("table", rows))
            continue

        if re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            flush_para()
            list_items = []
            ordered = bool(re.match(r"^\d+\.\s+", stripped))
            while i < len(lines) and (
                re.match(r"^[-*]\s+", lines[i].strip()) or re.match(r"^\d+\.\s+", lines[i].strip())
            ):
                item_text = re.sub(r"^([-*]|\d+\.)\s+", "", lines[i].strip())
                list_items.append(item_text)
                i += 1
            blocks.append(("list", (ordered, list_items)))
            continue

        para_buf.append(stripped)
        i += 1

    flush_para()
    return blocks


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[\[.+?\]\])")


def add_inline_runs(paragraph, text):
    """Handle **bold**, *italic*, `code`, and [[wikilink|Display]] -> Display (never leaked as raw
    brackets — the exact bug the 2026-07-22 export run had to fix by hand)."""
    for chunk in INLINE_RE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("*") and chunk.endswith("*") and not chunk.startswith("**"):
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
        elif chunk.startswith("[[") and chunk.endswith("]]"):
            inner = chunk[2:-2]
            display = inner.split("|")[-1]
            paragraph.add_run(display)
        else:
            paragraph.add_run(chunk)


# ---------------------------------------------------------------------------
# DOCX assembly
# ---------------------------------------------------------------------------

def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_toc_field(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and choose 'Update Field' to generate the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_text)
    run._r.append(fld_end)


def load_manifest():
    """Returns {slug: {path, title, type, section, geography, source_note}}."""
    manifest = {}
    if not os.path.exists(MANIFEST_PATH):
        return manifest
    with open(MANIFEST_PATH) as f:
        for line in f:
            if not line.startswith("| ") or line.startswith("| #") or line.startswith("|---"):
                continue
            cols = [c.strip() for c in line.strip().strip("|").split("|")]
            if len(cols) < 7:
                continue
            _, filename, title, type_, section, geography, source_note = cols[:7]
            filename = filename.strip("`")
            slug = os.path.splitext(filename)[0]
            manifest[slug] = dict(
                path=os.path.join(EXHIBITS_DIR, filename),
                title=title,
                type=type_,
                section=section,
                geography=geography,
                source_note=source_note,
            )
    return manifest


def render_blocks(document, blocks, manifest, figure_counter, smoke_test):
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            style = f"Heading {min(level, 3)}"
            p = document.add_heading(level=min(level, 3))
            p.style = document.styles[style]
            add_inline_runs(p, text)
        elif kind == "para":
            p = document.add_paragraph()
            add_inline_runs(p, payload)
        elif kind == "quote":
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(payload)
            run.italic = True
            run.font.color.rgb = DARK_GREY
        elif kind == "hr":
            p = document.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif kind == "list":
            ordered, items = payload
            style = "List Number" if ordered else "List Bullet"
            for item in items:
                p = document.add_paragraph(style=style)
                add_inline_runs(p, item)
        elif kind == "table":
            rows = payload
            if not rows:
                continue
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Light Grid Accent 2" if "Light Grid Accent 2" in [
                s.name for s in document.styles
            ] else "Table Grid"
            for r, row in enumerate(rows):
                for c, cell_text in enumerate(row):
                    if c >= len(table.columns):
                        continue
                    cell = table.cell(r, c)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline_runs(p, cell_text)
                    if r == 0:
                        for run in p.runs:
                            run.bold = True
            document.add_paragraph()
        elif kind == "figure":
            slug = payload
            entry = manifest.get(slug)
            if not entry or not os.path.exists(entry["path"]):
                p = document.add_paragraph()
                run = p.add_run(f"[Figure marker '{slug}' did not resolve to a manifest entry — flagged, not filled in.]")
                run.italic = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                continue
            figure_counter[0] += 1
            n = figure_counter[0]
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(entry["path"], width=Inches(5.5))
            cap = document.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(f"Figure {n}. {entry['title']}")
            cap_run.bold = True
            cap_run.font.size = Pt(9.5)
            src = document.add_paragraph()
            src.alignment = WD_ALIGN_PARAGRAPH.CENTER
            src_run = src.add_run(f"Source: {entry['source_note']}")
            src_run.italic = True
            src_run.font.size = Pt(8.5)
            src_run.font.color.rgb = DARK_GREY


def add_footer_page_numbers(document, watermark=None):
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if watermark:
        run = p.add_run(f"{watermark} — Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    else:
        p.add_run("Page ")
    add_page_number_field(p)


def get_tracker_title():
    if os.path.exists(TRACKER_PATH):
        with open(TRACKER_PATH) as f:
            for line in f:
                m = re.match(r"^#\s+(.*)$", line.strip())
                if m:
                    return m.group(1)
    return "talabat Group — AI-Enabled Capital Allocation and Performance Management for the 2026 USD175 Million Investment Programme"


def build(smoke_test):
    section_files = sorted(
        glob.glob(os.path.join(DRAFTS_DIR, "Section_*.md")),
        key=lambda p: int(re.search(r"Section_(\d+)", os.path.basename(p)).group(1)),
    )

    gate_failures = []
    for i in range(1, 15):
        matches = [p for p in section_files if int(re.search(r"Section_(\d+)", os.path.basename(p)).group(1)) == i]
        if not matches:
            gate_failures.append(f"Section {i}: does not exist under Business_Plan_Drafts_v2/")
            continue
        meta, _ = strip_frontmatter(open(matches[0]).read())
        status = meta.get("status", "")
        if "Done" not in status or "independently verified" not in status:
            gate_failures.append(f"Section {i}: status is '{status[:80]}...' — not Done (independently verified)")

    if not smoke_test:
        gate_failures.extend(editorial_preflight(section_files))

    if gate_failures and not smoke_test:
        print("EXPORT BLOCKED — input gate failed (Publication_Layer.md §2):")
        for f in gate_failures:
            print(f"  - {f}")
        print("\nRun with --smoke-test to verify the export pipeline mechanically works anyway "
              "(output is watermarked and written outside Outputs/, never a real submission).")
        sys.exit(1)

    watermark = "CAPABILITY VERIFICATION ONLY — NOT A COMPLETE OR INDEPENDENTLY VERIFIED BUSINESS PLAN" if smoke_test else None

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Cover page ---
    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(get_tracker_title())
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = ORANGE

    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("AI Business Plan — GSB Template v2.0 (McKinsey Edition)")
    subtitle_run.font.size = Pt(14)

    meta_p = document.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(
        "AASTMT MBA, AI for Business Organizations Track — Group G02\n"
        "Prepared by the Agentic OS — draft, internal working copy"
    )

    if watermark:
        warn_p = document.add_paragraph()
        warn_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        warn_run = warn_p.add_run(watermark)
        warn_run.bold = True
        warn_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        warn_run.font.size = Pt(12)

    document.add_page_break()

    # --- TOC page ---
    toc_heading = document.add_heading("Table of Contents", level=1)
    add_toc_field(document)
    document.add_page_break()

    manifest = load_manifest()
    figure_counter = [0]

    for path in section_files:
        text = open(path).read()
        meta, body = strip_frontmatter(text)
        blocks = parse_blocks(body)
        render_blocks(document, blocks, manifest, figure_counter, smoke_test)
        document.add_page_break()

    add_footer_page_numbers(document, watermark=watermark)

    out_dir = SMOKE_TEST_DIR if smoke_test else OUTPUTS_DIR
    os.makedirs(out_dir, exist_ok=True)
    basename = "SMOKE_TEST_Talabat_Group_AI_Investment_Allocation_Business_Plan" if smoke_test else "Talabat_Group_AI_Investment_Allocation_Business_Plan"
    docx_path = os.path.join(out_dir, f"{basename}.docx")
    document.save(docx_path)
    print(f"Wrote {docx_path}")

    # PDF: rendered from an independently-built HTML document (same content, same figures), printed
    # via headless Chromium — see module docstring for why this replaces the documented-but-broken
    # LibreOffice DOCX->PDF path in this environment.
    title = get_tracker_title()
    html = render_full_html(section_files, manifest, smoke_test, title)
    html_path = os.path.join(out_dir, f"{basename}.html")
    with open(html_path, "w") as f:
        f.write(html)
    pdf_path = os.path.join(out_dir, f"{basename}.pdf")
    pdf_path = convert_to_pdf_via_chromium(html_path, pdf_path, watermark)

    return docx_path, pdf_path, gate_failures, figure_counter[0], len(section_files)


def find_chromium():
    """Locate a usable Chromium executable. Prefers this repo's documented pre-installed browser
    (PLAYWRIGHT_BROWSERS_PATH=/opt/pw-browsers); falls back to Playwright's own default resolution
    if that path layout ever changes."""
    base = os.environ.get("PLAYWRIGHT_BROWSERS_PATH", "/opt/pw-browsers")
    if os.path.isdir(base):
        for entry in sorted(os.listdir(base)):
            if entry.startswith("chromium-"):
                candidate = os.path.join(base, entry, "chrome-linux", "chrome")
                if os.path.exists(candidate):
                    return candidate
    return None


HTML_CSS = """
body { font-family: Calibri, Arial, sans-serif; color: #222; font-size: 11pt; line-height: 1.4; }
h1 { color: #FF6B00; font-size: 18pt; page-break-before: always; }
h1:first-of-type { page-break-before: avoid; }
h2 { font-size: 14pt; margin-top: 1.2em; }
h3 { font-size: 12pt; margin-top: 1em; }
table { border-collapse: collapse; width: 100%; margin: 0.8em 0; font-size: 9.5pt; }
th, td { border: 1px solid #ccc; padding: 4px 6px; text-align: left; vertical-align: top; }
th { background: #f2f2f2; font-weight: bold; }
blockquote { border-left: 3px solid #ccc; margin: 0.6em 0; padding-left: 0.8em; color: #555; font-style: italic; font-size: 9.5pt; }
hr { border: none; border-top: 1px solid #ccc; margin: 1em 0; }
.figure { text-align: center; margin: 1em 0; }
.figure img { max-width: 90%; }
.figure-caption { font-weight: bold; font-size: 9.5pt; }
.figure-source { font-style: italic; font-size: 8.5pt; color: #555; }
.cover { text-align: center; page-break-after: always; padding-top: 3in; }
.cover h1 { page-break-before: avoid; font-size: 22pt; }
.cover .subtitle { font-size: 14pt; }
.watermark { color: #cc0000; font-weight: bold; font-size: 12pt; }
.toc { page-break-after: always; }
"""


def blocks_to_html(blocks, manifest, figure_counter):
    def esc(t):
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def inline_html(text):
        out = []
        for chunk in INLINE_RE.split(text):
            if not chunk:
                continue
            if chunk.startswith("**") and chunk.endswith("**"):
                out.append(f"<strong>{esc(chunk[2:-2])}</strong>")
            elif chunk.startswith("*") and chunk.endswith("*") and not chunk.startswith("**"):
                out.append(f"<em>{esc(chunk[1:-1])}</em>")
            elif chunk.startswith("`") and chunk.endswith("`"):
                out.append(f"<code>{esc(chunk[1:-1])}</code>")
            elif chunk.startswith("[[") and chunk.endswith("]]"):
                out.append(esc(chunk[2:-2].split("|")[-1]))
            else:
                out.append(esc(chunk))
        return "".join(out)

    html = []
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            tag = f"h{min(level, 3)}"
            html.append(f"<{tag}>{inline_html(text)}</{tag}>")
        elif kind == "para":
            html.append(f"<p>{inline_html(payload)}</p>")
        elif kind == "quote":
            html.append(f"<blockquote>{inline_html(payload)}</blockquote>")
        elif kind == "hr":
            html.append("<hr>")
        elif kind == "list":
            ordered, items = payload
            tag = "ol" if ordered else "ul"
            items_html = "".join(f"<li>{inline_html(i)}</li>" for i in items)
            html.append(f"<{tag}>{items_html}</{tag}>")
        elif kind == "table":
            rows = payload
            if not rows:
                continue
            rows_html = []
            header, *body_rows = rows
            rows_html.append("<tr>" + "".join(f"<th>{inline_html(c)}</th>" for c in header) + "</tr>")
            for row in body_rows:
                rows_html.append("<tr>" + "".join(f"<td>{inline_html(c)}</td>" for c in row) + "</tr>")
            html.append("<table>" + "".join(rows_html) + "</table>")
        elif kind == "figure":
            slug = payload
            entry = manifest.get(slug)
            if not entry or not os.path.exists(entry["path"]):
                html.append(f"<p style='color:#cc0000'><em>[Figure marker '{slug}' did not resolve.]</em></p>")
                continue
            figure_counter[0] += 1
            n = figure_counter[0]
            html.append(
                f"<div class='figure'><img src='file://{entry['path']}'>"
                f"<div class='figure-caption'>Figure {n}. {esc(entry['title'])}</div>"
                f"<div class='figure-source'>Source: {esc(entry['source_note'])}</div></div>"
            )
    return "\n".join(html)


def render_full_html(section_files, manifest, smoke_test, title):
    figure_counter = [0]
    body_parts = []
    for path in section_files:
        text = open(path).read()
        _, body = strip_frontmatter(text)
        blocks = parse_blocks(body)
        body_parts.append(blocks_to_html(blocks, manifest, figure_counter))

    watermark_html = ""
    if smoke_test:
        watermark_html = (
            "<div class='watermark'>CAPABILITY VERIFICATION ONLY — "
            "NOT A COMPLETE OR INDEPENDENTLY VERIFIED BUSINESS PLAN</div>"
        )

    cover = (
        f"<div class='cover'><h1>{title}</h1>"
        "<div class='subtitle'>AI Business Plan — GSB Template v2.0 (McKinsey Edition)</div>"
        "<p>AASTMT MBA, AI for Business Organizations Track — Group G02<br>"
        "Prepared by the Agentic OS — draft, internal working copy</p>"
        f"{watermark_html}</div>"
    )

    toc_items = []
    for path in section_files:
        text = open(path).read()
        meta, _ = strip_frontmatter(text)
        toc_items.append(f"<li>Section {meta.get('section', '?')}. {meta.get('title', os.path.basename(path))}</li>")
    toc = f"<div class='toc'><h1>Table of Contents</h1><ol>{''.join(toc_items)}</ol></div>"

    return f"<html><head><meta charset='utf-8'><style>{HTML_CSS}</style></head><body>{cover}{toc}{''.join(body_parts)}</body></html>"


def convert_to_pdf_via_chromium(html_path, pdf_path, watermark):
    chromium = find_chromium()
    if not chromium:
        print("PDF conversion FAILED: no Chromium executable found under PLAYWRIGHT_BROWSERS_PATH.")
        return None
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("PDF conversion FAILED: playwright not installed (pip install playwright).")
        return None

    footer_note = f"{watermark} — " if watermark else ""
    with sync_playwright() as p:
        browser = p.chromium.launch(executable_path=chromium)
        page = browser.new_page()
        page.goto(f"file://{html_path}")
        page.pdf(
            path=pdf_path,
            format="A4",
            print_background=True,
            display_header_footer=True,
            header_template="<span></span>",
            footer_template=(
                f"<div style='font-size:8px; width:100%; text-align:center; color:#888;'>"
                f"{footer_note}Page <span class='pageNumber'></span> of <span class='totalPages'></span></div>"
            ),
            margin={"top": "0.7in", "bottom": "0.7in", "left": "0.7in", "right": "0.7in"},
        )
        browser.close()
    if not os.path.exists(pdf_path):
        print("PDF conversion FAILED: Playwright ran but no PDF was written.")
        return None
    print(f"Wrote {pdf_path}")
    return pdf_path


def estimate_page_count(pdf_path):
    """Best-effort page count via a raw byte-pattern count (no reliable PDF library available in
    this environment) — counts distinct page objects. Report as an estimate, not authoritative."""
    if not pdf_path or not os.path.exists(pdf_path):
        return None
    with open(pdf_path, "rb") as f:
        data = f.read()
    count = len(re.findall(rb"/Type\s*/Page[^s]", data))
    return count if count > 0 else None


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--smoke-test", action="store_true", help="Bypass the gate, watermark output, write outside Outputs/")
    args = parser.parse_args()

    docx_path, pdf_path, gate_failures, n_figures, n_sections = build(args.smoke_test)
    pages = estimate_page_count(pdf_path)

    print("\n--- Run summary ---")
    print(f"Sections assembled: {n_sections}/14")
    print(f"Figures embedded: {n_figures}")
    print(f"Gate failures: {len(gate_failures)}")
    print(f"DOCX: {docx_path} ({os.path.getsize(docx_path)} bytes)")
    if pdf_path:
        print(f"PDF: {pdf_path} ({os.path.getsize(pdf_path)} bytes), ~{pages} pages (estimate)")
    else:
        print("PDF: not produced")
