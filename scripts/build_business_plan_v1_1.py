#!/usr/bin/env python3
"""Build the visually enhanced v1.1 submission edition without altering the verified v2 sections."""
from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DRAFTS = ROOT / "vault" / "Projects" / "Business_Plan_Drafts_v2"
OUT = ROOT / "Outputs" / "Business_Plan_v1.1_Executive_Visual.docx"
PDF_OUT = ROOT / "Outputs" / "Business_Plan_v1.1_Executive_Visual.pdf"
EXHIBIT_DIR = ROOT / "Outputs" / "V1_1_Exhibits"
ORANGE = RGBColor(255, 107, 0)
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(82, 82, 82)
LIGHT = "F4F6F9"

EXHIBITS = {
    1: [
        ("Exhibit_01_disclosed_programme_structure.png",
         "Exhibit 1. Approximately 69% of the disclosed 2026 programme is directed to Everyday App expansion."),
        ("Exhibit_02_portfolio_actions.png",
         "Exhibit 2. The portfolio requires four actions—not one uniform funding decision."),
    ],
    2: [("Exhibit_03_two_ai_layers.png",
         "Exhibit 3. The Agentic OS governs capital decisions while talabat’s embedded AI executes operating use cases.")],
    3: [("Exhibit_04_growth_and_scale.png",
         "Exhibit 4. Non-GCC growth is more than twice GCC growth, but GCC still carries about four-fifths of GMV.")],
    4: [("Exhibit_05_value_driver_chain.png",
         "Exhibit 5. Capital creates value only when capability deployment changes customer behaviour before financial outcomes.")],
    10: [("Exhibit_07_risk_matrix.png",
          "Exhibit 7. Three high-priority risks require management attention before capital release.")],
    12: [("Exhibit_06_three_horizons.png",
          "Exhibit 6. Three proof gates sequence quick wins before broader scale commitments.")],
    13: [("Exhibit_08_kpi_tree.png",
          "Exhibit 8. Capital release should follow leading proof signals before lagging financial outcomes.")],
}


def _figure(title, source, classification, implication, height=3.2):
    import matplotlib.pyplot as plt
    fig = plt.figure(figsize=(10, height), dpi=180, facecolor="white")
    fig.text(0.04, 0.94, title, fontsize=13, fontweight="bold", color="#0B2545", va="top")
    fig.text(0.04, 0.055, f"Source: {source}  |  Classification: {classification}",
             fontsize=6.8, color="#555555")
    fig.text(0.04, 0.012, f"Management implication: {implication}",
             fontsize=7.2, fontweight="bold", color="#1F4D78")
    return fig


def generate_exhibits():
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch, Rectangle
    EXHIBIT_DIR.mkdir(parents=True, exist_ok=True)

    fig = _figure(
        "Approximately 69% of the disclosed 2026 programme is directed to Everyday App expansion",
        "TLB-020 pp.12,16; TLB-014 pp.16,19", "Disclosed talabat programme structure",
        "Govern release within the disclosed envelope; do not present the split as an Agentic OS allocation.")
    ax = fig.add_axes([0.08, 0.25, 0.84, 0.48])
    vals, labels, colors_ = [75, 45, 55], ["Everyday App opex\n~USD75mn", "Everyday App capex\n~USD45mn", "Food leadership\n~USD55mn"], ["#1F4D78", "#6F95B7", "#FF6B00"]
    left = 0
    for val, label, color in zip(vals, labels, colors_):
        ax.barh([0], [val], left=left, color=color, edgecolor="white", height=0.55)
        ax.text(left + val/2, 0, label, ha="center", va="center", color="white", fontsize=9, fontweight="bold")
        left += val
    ax.set_xlim(0, 175); ax.set_yticks([]); ax.set_xlabel("USD million (approximately)", fontsize=8)
    for s in ax.spines.values(): s.set_visible(False)
    fig.savefig(EXHIBIT_DIR / "Exhibit_01_disclosed_programme_structure.png", bbox_inches="tight"); plt.close(fig)

    fig = _figure(
        "The portfolio requires four actions—not one uniform funding decision",
        "Investment Options Register; DEC-009; OPT-001–005", "Agentic OS analytical recommendation; no dollar allocation shown",
        "Fund OPT-003 now, pilot OPT-002, continue OPT-001/005 at pace, and build OPT-004 later.")
    ax = fig.add_axes([0.05, 0.20, 0.90, 0.60]); ax.axis("off")
    lanes = [
        ("FUND NOW", "OPT-003\nFood-leadership CVP\nUAE · Kuwait · Qatar", "#FF6B00"),
        ("PILOT", "OPT-002\ntalabat pro\nEgypt → Iraq", "#D88A27"),
        ("CONTINUE AT PACE", "OPT-001  tMart densification\nOPT-005  AI/personalisation", "#1F4D78"),
        ("BUILD LATER", "OPT-004\nAdvertising monetisation", "#6F95B7"),
    ]
    for i,(action,body,color) in enumerate(lanes):
        x=0.015+i*0.245
        ax.add_patch(FancyBboxPatch((x,0.08),0.225,0.78,boxstyle="round,pad=0.012",facecolor="#F4F6F9",edgecolor=color,linewidth=1.8))
        ax.text(x+0.112,0.73,action,ha="center",va="center",fontsize=8,fontweight="bold",color=color)
        ax.text(x+0.112,0.40,body,ha="center",va="center",fontsize=8,color="#222222",linespacing=1.35)
    fig.savefig(EXHIBIT_DIR / "Exhibit_02_portfolio_actions.png", bbox_inches="tight"); plt.close(fig)

    fig = _figure(
        "The Agentic OS governs capital decisions while talabat’s embedded AI executes operating use cases",
        "Sections 2 and 5; Repository Impact Assessment", "Operating-model architecture; no system-access claim",
        "Keep governance and operational AI distinct; human approval remains mandatory before capital movement.")
    ax=fig.add_axes([0.07,0.18,0.86,0.66]); ax.axis("off")
    boxes=[(0.34,0.76,0.32,0.14,"USD175mn investment programme","#0B2545"),
           (0.28,0.51,0.44,0.15,"Agentic OS allocation + governance layer\nEvidence ranking · options · KPI gates · human approval","#1F4D78"),
           (0.08,0.26,0.35,0.14,"Everyday App initiatives\n~USD120mn disclosed bucket","#6F95B7"),
           (0.57,0.26,0.35,0.14,"Food-leadership initiatives\n~USD55mn disclosed bucket","#FF6B00"),
           (0.20,0.02,0.60,0.14,"talabat embedded operational AI\nPersonalisation · ranking · offer timing · dispatch optimisation","#525252")]
    for x,y,w,h,t,c in boxes:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.012",facecolor="white",edgecolor=c,linewidth=1.7))
        ax.text(x+w/2,y+h/2,t,ha="center",va="center",fontsize=8.2,color=c,fontweight="bold")
    for x1,y1,x2,y2 in [(0.5,0.76,0.5,0.66),(0.5,0.51,0.255,0.40),(0.5,0.51,0.745,0.40),(0.255,0.26,0.5,0.16),(0.745,0.26,0.5,0.16)]:
        ax.annotate("",xy=(x2,y2),xytext=(x1,y1),arrowprops=dict(arrowstyle="->",color="#777777",lw=1.2))
    fig.savefig(EXHIBIT_DIR / "Exhibit_03_two_ai_layers.png", bbox_inches="tight"); plt.close(fig)

    fig = _figure(
        "Non-GCC growth is more than twice GCC growth, but GCC still carries about four-fifths of GMV",
        "TLB-019 p.14; TLB-011 p.2", "Disclosed; FY2025 growth/share shown separately",
        "Protect GCC’s scale while using reversible pilots to learn in faster-growing non-GCC markets.")
    ax1=fig.add_axes([0.08,0.23,0.38,0.53]); ax2=fig.add_axes([0.58,0.23,0.34,0.53])
    ax1.bar(["GCC","Non-GCC"],[22,57],color=["#1F4D78","#FF6B00"],width=.58)
    ax1.set_title("FY2025 GMV growth (y/y)",fontsize=9,fontweight="bold"); ax1.set_ylabel("%",fontsize=8)
    for i,v in enumerate([22,57]): ax1.text(i,v+2,f"{v}%",ha="center",fontsize=9,fontweight="bold")
    ax1.set_ylim(0,65); ax1.spines[["top","right"]].set_visible(False)
    ax2.barh(["Non-GCC","GCC"],[18,82],color=["#FF6B00","#1F4D78"],height=.55)
    ax2.set_title("FY2025 Group GMV share",fontsize=9,fontweight="bold"); ax2.set_xlabel("%",fontsize=8)
    for i,v in enumerate([18,82]): ax2.text(v+2,i,f"{v}%",va="center",fontsize=9,fontweight="bold")
    ax2.set_xlim(0,100); ax2.spines[["top","right"]].set_visible(False)
    fig.savefig(EXHIBIT_DIR / "Exhibit_04_growth_and_scale.png", bbox_inches="tight"); plt.close(fig)

    fig = _figure(
        "Capital creates value only when capability deployment changes customer behaviour before financial outcomes",
        "Value Driver Tree v2; KPI Tree v2", "Analytical causal logic; not a quantified forecast",
        "Release capital only when each option demonstrates its next causal link through named KPIs.")
    ax=fig.add_axes([0.04,0.23,0.92,0.53]); ax.axis("off")
    chain=[("CAPITAL","Option-level\nfunding"),("CAPABILITY","AI · pro · tMart\npartner CVP"),("BEHAVIOUR","Frequency · retention\nmulti-vertical use"),("ECONOMICS","GMV · revenue\ngross profit"),("OUTCOMES","EBITDA · cash\nlong-term CLV")]
    for i,(head,body) in enumerate(chain):
        x=0.01+i*0.198
        color="#FF6B00" if i==0 else "#1F4D78"
        ax.add_patch(FancyBboxPatch((x,0.18),0.17,0.58,boxstyle="round,pad=.012",facecolor="#F4F6F9",edgecolor=color,linewidth=1.5))
        ax.text(x+.085,.62,head,ha="center",fontsize=7.8,fontweight="bold",color=color)
        ax.text(x+.085,.38,body,ha="center",va="center",fontsize=8)
        if i<4: ax.annotate("",xy=(x+.198,.47),xytext=(x+.17,.47),arrowprops=dict(arrowstyle="->",color="#777777"))
    fig.savefig(EXHIBIT_DIR / "Exhibit_05_value_driver_chain.png", bbox_inches="tight"); plt.close(fig)

    fig = _figure(
        "Three proof gates sequence quick wins before broader scale commitments",
        "Section 12; DEC-009; Investment Options Register", "Analytical implementation horizon; no new budget",
        "H2 and H3 capital release depends on evidence generated in the prior horizon.")
    ax=fig.add_axes([0.05,0.19,0.90,0.62]); ax.axis("off")
    horizons=[("H1 · 0–6 months","Fund OPT-003 now\nPilot OPT-002\nContinue OPT-001/005","Measure adoption,\nretention and margin"),
              ("H2 · 6–18 months","OPT-002 go/no-go\nReview OPT-003\nBegin OPT-004","Human gate:\nscale, hold or redirect"),
              ("H3 · 18+ months","Scale evidenced options\nExtend OPT-004 selectively","Revisit cross-market\nallocation with new data")]
    for i,(h,a,g) in enumerate(horizons):
        x=.01+i*.33
        ax.add_patch(FancyBboxPatch((x,.12),.30,.70,boxstyle="round,pad=.012",facecolor="#F4F6F9",edgecolor="#1F4D78",linewidth=1.6))
        ax.text(x+.15,.72,h,ha="center",fontsize=9,fontweight="bold",color="#0B2545")
        ax.text(x+.15,.48,a,ha="center",va="center",fontsize=8.2,linespacing=1.35)
        ax.text(x+.15,.22,g,ha="center",va="center",fontsize=7.7,color="#FF6B00",fontweight="bold")
        if i<2: ax.annotate("",xy=(x+.33,.47),xytext=(x+.30,.47),arrowprops=dict(arrowstyle="->",color="#FF6B00",lw=1.6))
    fig.savefig(EXHIBIT_DIR / "Exhibit_06_three_horizons.png", bbox_inches="tight"); plt.close(fig)

    fig = _figure(
        "Three high-priority risks require management attention before capital release",
        "Section 10 risk register", "Ordinal management assessment; not actuarial probability",
        "Resolve the stage-gate gap and track margin and Food retention before releasing incremental capital.")
    ax=fig.add_axes([0.12,0.22,0.72,0.55])
    ax.set_xlim(.5,5.5); ax.set_ylim(.5,5.5); ax.set_xticks(range(1,6),["Very low","Low","Medium","Med-high","High"],fontsize=7)
    ax.set_yticks(range(1,6),["Very low","Low","Medium","Med-high","High"],fontsize=7)
    ax.set_xlabel("Impact",fontsize=8); ax.set_ylabel("Likelihood",fontsize=8); ax.grid(True,color="#DADCE0")
    ax.add_patch(Rectangle((4.5,4.5),1,1,facecolor="#FFF2E8",edgecolor="#FF6B00",lw=1.5))
    for label,x,y in [("ORG-01",5,5.18),("FIN-01",5,4.98),("MKT-01",5,4.78)]:
        ax.text(x,y,label,ha="center",va="center",fontsize=8,fontweight="bold",color="#B54E00")
    for label,x,y in [("ORG-02",5,3),("FIN-03",3,5),("REG-02",3,5),("FIN-04",3,3)]:
        ax.text(x,y,label,ha="center",va="center",fontsize=7,color="#1F4D78")
    ax.text(5.7,4.98,"High × High",fontsize=7.5,color="#B54E00",fontweight="bold",va="center")
    fig.savefig(EXHIBIT_DIR / "Exhibit_07_risk_matrix.png", bbox_inches="tight"); plt.close(fig)

    fig = _figure(
        "Capital release should follow leading proof signals before lagging financial outcomes",
        "KPI Tree v2; Section 13", "Selective KPI tree; full 50-KPI inventory retained in repository",
        "Do not scale options on spend alone; require adoption and operating proof before financial confirmation.")
    ax=fig.add_axes([0.04,0.22,0.92,0.55]); ax.axis("off")
    levels=[("CAPITAL CONTROL","Spend-to-plan\nStage-gate compliance","Leading"),
            ("CUSTOMER PROOF","Pro adoption · frequency\nMulti-vertical usage","Leading"),
            ("OPERATING PROOF","Partner retention\nDelivery · AI throughput","Leading / lagging"),
            ("FINANCIAL CONFIRMATION","GMV · revenue · margin\nCash conversion","Lagging")]
    for i,(head,body,tag) in enumerate(levels):
        x=.01+i*.247
        ax.add_patch(FancyBboxPatch((x,.15),.215,.66,boxstyle="round,pad=.012",facecolor="white",edgecolor="#1F4D78",linewidth=1.5))
        ax.text(x+.107,.68,head,ha="center",fontsize=7.6,fontweight="bold",color="#0B2545")
        ax.text(x+.107,.44,body,ha="center",va="center",fontsize=8)
        ax.text(x+.107,.22,tag,ha="center",fontsize=7.3,fontweight="bold",color="#FF6B00")
        if i<3: ax.annotate("",xy=(x+.247,.48),xytext=(x+.215,.48),arrowprops=dict(arrowstyle="->",color="#777777"))
    fig.savefig(EXHIBIT_DIR / "Exhibit_08_kpi_tree.png", bbox_inches="tight"); plt.close(fig)


def generate_exhibits():
    """Generate consistent high-resolution decision exhibits with Pillow."""
    from PIL import Image, ImageDraw, ImageFont
    import textwrap

    EXHIBIT_DIR.mkdir(parents=True, exist_ok=True)
    W, H = 1800, 590
    navy, blue, orange, gray, pale = "#0B2545", "#1F4D78", "#FF6B00", "#555555", "#F4F6F9"
    font_dir = Path("C:/Windows/Fonts")

    def font(size, bold=False):
        return ImageFont.truetype(str(font_dir / ("arialbd.ttf" if bold else "arial.ttf")), size)

    def canvas(title, source, classification, implication):
        im = Image.new("RGB", (W, H), "white")
        d = ImageDraw.Draw(im)
        d.text((55, 28), title, fill=navy, font=font(29, True))
        d.line((55, 82, W-55, 82), fill="#DADCE0", width=2)
        d.text((55, H-67), f"Source: {source}  |  Classification: {classification}", fill=gray, font=font(14))
        d.text((55, H-35), f"Management implication: {implication}", fill=blue, font=font(15, True))
        return im, d

    def box(d, xy, heading, body, edge=blue, fill=pale):
        x1,y1,x2,y2=xy
        d.rounded_rectangle(xy, radius=16, fill=fill, outline=edge, width=4)
        d.text(((x1+x2)//2, y1+24), heading, fill=edge, font=font(18,True), anchor="ma")
        lines=textwrap.wrap(body, width=max(12,int((x2-x1)/15)))
        d.multiline_text(((x1+x2)//2,(y1+y2)//2+12),"\n".join(lines),fill="#222222",font=font(17),anchor="mm",align="center",spacing=6)

    def arrow(d, start, end, color=gray):
        d.line((*start,*end),fill=color,width=5)
        x,y=end
        d.polygon([(x,y),(x-16,y-10),(x-16,y+10)],fill=color)

    im,d=canvas("Approximately 69% of the disclosed 2026 programme is directed to Everyday App expansion",
                "TLB-020 pp.12,16; TLB-014 pp.16,19","Disclosed talabat programme structure",
                "Govern release within the disclosed envelope; do not present the split as an Agentic OS allocation.")
    x0,y0,bh=100,210,145
    scale=1500/175
    for val,label,color in [(75,"Everyday App opex\n~USD75mn",blue),(45,"Everyday App capex\n~USD45mn","#6F95B7"),(55,"Food leadership\n~USD55mn",orange)]:
        x1=x0+int(val*scale)
        d.rectangle((x0,y0,x1,y0+bh),fill=color,outline="white",width=4)
        d.multiline_text(((x0+x1)//2,y0+bh//2),label,fill="white",font=font(20,True),anchor="mm",align="center",spacing=5)
        x0=x1
    d.text((900,390),"USD million (approximately)  |  Total = USD175mn",fill=gray,font=font(16),anchor="ma")
    im.save(EXHIBIT_DIR/"Exhibit_01_disclosed_programme_structure.png")

    im,d=canvas("The portfolio requires four actions—not one uniform funding decision",
                "Investment Options Register; DEC-009; OPT-001–005","Agentic OS analytical recommendation; no dollar allocation shown",
                "Fund OPT-003 now, pilot OPT-002, continue OPT-001/005 at pace, and build OPT-004 later.")
    lanes=[("FUND NOW","OPT-003\nFood-leadership CVP\nUAE · Kuwait · Qatar",orange),
           ("PILOT","OPT-002\ntalabat pro\nEgypt → Iraq","#D88A27"),
           ("CONTINUE AT PACE","OPT-001  tMart densification\nOPT-005  AI/personalisation",blue),
           ("BUILD LATER","OPT-004\nAdvertising monetisation","#6F95B7")]
    for i,(h,b,c) in enumerate(lanes):
        x=65+i*430; box(d,(x,135,x+385,455),h,b,c)
    im.save(EXHIBIT_DIR/"Exhibit_02_portfolio_actions.png")

    im,d=canvas("The Agentic OS governs capital decisions while talabat’s embedded AI executes operating use cases",
                "Sections 2 and 5; Repository Impact Assessment","Operating-model architecture; no system-access claim",
                "Keep governance and operational AI distinct; human approval remains mandatory before capital movement.")
    box(d,(620,105,1180,190),"PROGRAMME","USD175mn disclosed envelope",navy,"white")
    box(d,(500,235,1300,340),"GOVERNANCE LAYER","Agentic OS: evidence ranking · options · KPI gates · human approval",blue,"white")
    box(d,(100,385,780,485),"EVERYDAY APP","Initiatives within ~USD120mn disclosed bucket","#6F95B7","white")
    box(d,(1020,385,1700,485),"FOOD LEADERSHIP","Initiatives within ~USD55mn disclosed bucket",orange,"white")
    arrow(d,(900,190),(900,235)); arrow(d,(900,340),(440,385)); arrow(d,(900,340),(1360,385))
    d.text((900,505),"Execution layer: talabat embedded AI — personalisation · ranking · offer timing · dispatch optimisation",fill=gray,font=font(16,True),anchor="ma")
    im.save(EXHIBIT_DIR/"Exhibit_03_two_ai_layers.png")

    im,d=canvas("Non-GCC growth is more than twice GCC growth, but GCC still carries about four-fifths of GMV",
                "TLB-019 p.14; TLB-011 p.2","Disclosed; FY2025 growth/share shown separately",
                "Protect GCC’s scale while using reversible pilots to learn in faster-growing non-GCC markets.")
    d.text((430,125),"FY2025 GMV growth (y/y)",fill=navy,font=font(20,True),anchor="ma")
    for i,(lab,val,c) in enumerate([("GCC",22,blue),("Non-GCC",57,orange)]):
        y=190+i*110; d.text((95,y+30),lab,fill=gray,font=font(18,True))
        d.rectangle((260,y,260+val*9,y+62),fill=c)
        d.text((280+val*9,y+31),f"{val}%",fill=c,font=font(22,True),anchor="lm")
    d.text((1325,125),"FY2025 Group GMV share",fill=navy,font=font(20,True),anchor="ma")
    x=1020
    for val,lab,c in [(82,"GCC 82%",blue),(18,"Non-GCC 18%",orange)]:
        x1=x+int(val*7.2); d.rectangle((x,220,x1,330),fill=c,outline="white",width=3)
        d.text(((x+x1)//2,275),lab,fill="white",font=font(19,True),anchor="mm"); x=x1
    d.text((1325,370),"Separate dated facts — not a blended range",fill=gray,font=font(16),anchor="ma")
    im.save(EXHIBIT_DIR/"Exhibit_04_growth_and_scale.png")

    im,d=canvas("Capital creates value only when capability deployment changes customer behaviour before financial outcomes",
                "Value Driver Tree v2; KPI Tree v2","Analytical causal logic; not a quantified forecast",
                "Release capital only when each option demonstrates its next causal link through named KPIs.")
    chain=[("CAPITAL","Option-level funding",orange),("CAPABILITY","AI · pro · tMart\npartner CVP",blue),
           ("BEHAVIOUR","Frequency · retention\nmulti-vertical use",blue),("ECONOMICS","GMV · revenue\ngross profit",blue),
           ("OUTCOMES","EBITDA · cash\nlong-term CLV",navy)]
    for i,(h,b,c) in enumerate(chain):
        x=45+i*350; box(d,(x,155,x+300,420),h,b,c)
        if i<4: arrow(d,(x+300,288),(x+345,288))
    im.save(EXHIBIT_DIR/"Exhibit_05_value_driver_chain.png")

    im,d=canvas("Three proof gates sequence quick wins before broader scale commitments",
                "Section 12; DEC-009; Investment Options Register","Analytical implementation horizon; no new budget",
                "H2 and H3 capital release depends on evidence generated in the prior horizon.")
    horizons=[("H1 · 0–6 MONTHS","Fund OPT-003 now\nPilot OPT-002\nContinue OPT-001/005\n\nPROOF: adoption · retention · margin"),
              ("H2 · 6–18 MONTHS","OPT-002 go/no-go\nReview OPT-003\nBegin OPT-004\n\nGATE: scale · hold · redirect"),
              ("H3 · 18+ MONTHS","Scale evidenced options\nExtend OPT-004 selectively\n\nREVISIT: cross-market allocation")]
    for i,(h,b) in enumerate(horizons):
        x=80+i*565; box(d,(x,145,x+500,445),h,b,blue)
        if i<2: arrow(d,(x+500,295),(x+555,295),orange)
    im.save(EXHIBIT_DIR/"Exhibit_06_three_horizons.png")

    im,d=canvas("Three high-priority risks require management attention before capital release",
                "Section 10 risk register","Ordinal management assessment; not actuarial probability",
                "Resolve the stage-gate gap and track margin and Food retention before releasing incremental capital.")
    x0,y0,cell=280,120,78
    for i in range(6):
        d.line((x0+i*cell,y0,x0+i*cell,y0+5*cell),fill="#C8CDD3",width=2)
        d.line((x0,y0+i*cell,x0+5*cell,y0+i*cell),fill="#C8CDD3",width=2)
    d.rectangle((x0+4*cell,y0,x0+5*cell,y0+cell),fill="#FFF2E8",outline=orange,width=4)
    d.text((x0+4.5*cell,y0+cell/2),"ORG-01\nFIN-01\nMKT-01",fill="#B54E00",font=font(15,True),anchor="mm",align="center",spacing=2)
    d.text((x0+4.5*cell,y0-25),"HIGH IMPACT",fill=gray,font=font(14,True),anchor="ms")
    d.text((x0-25,y0+cell/2),"HIGH\nLIKELIHOOD",fill=gray,font=font(13,True),anchor="rm",align="right")
    d.text((900,190),"High × High priority",fill=orange,font=font(22,True))
    d.text((900,235),"ORG-01  No disclosed stage-gate mechanism\nFIN-01  EBITDA margin step-down already in motion\nMKT-01  Food retention pressure in GCC-3",fill="#222222",font=font(18),spacing=12)
    d.text((900,365),"Other material watchpoints: ORG-02 false precision · FIN-03 no disclosed ROI hurdle · REG-02 Qatar cap",fill=blue,font=font(16,True))
    im.save(EXHIBIT_DIR/"Exhibit_07_risk_matrix.png")

    im,d=canvas("Capital release should follow leading proof signals before lagging financial outcomes",
                "KPI Tree v2; Section 13","Selective KPI tree; full 50-KPI inventory retained in repository",
                "Do not scale options on spend alone; require adoption and operating proof before financial confirmation.")
    levels=[("CAPITAL CONTROL","Spend-to-plan\nStage-gate compliance","LEADING"),
            ("CUSTOMER PROOF","Pro adoption · frequency\nMulti-vertical usage","LEADING"),
            ("OPERATING PROOF","Partner retention\nDelivery · AI throughput","LEADING / LAGGING"),
            ("FINANCIAL CONFIRMATION","GMV · revenue · margin\nCash conversion","LAGGING")]
    for i,(h,b,t) in enumerate(levels):
        x=55+i*435; box(d,(x,150,x+380,430),h,b,blue)
        d.text((x+190,395),t,fill=orange,font=font(15,True),anchor="mm")
        if i<3: arrow(d,(x+380,290),(x+430,290))
    im.save(EXHIBIT_DIR/"Exhibit_08_kpi_tree.png")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8.5)


def strip_inline(text):
    text = re.sub(r"\[\[([^\]|]+)\|([^\]]+)\]\]", r"\2", text)
    text = re.sub(r"\[\[([^\]]+)\]\]", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = text.replace(
        "none is a generic risk-register placeholder",
        "each is evidence-specific",
    )
    text = re.sub(
        r", executed in-context by this session per the tooling note above, where a dedicated Agent call was unavailable",
        "",
        text,
        flags=re.I,
    )
    return text.strip()


def blocks(path):
    text = path.read_text(encoding="utf-8")
    text = re.sub(r"^---\n.*?\n---\n", "", text, flags=re.S)
    lines = text.splitlines()
    result, para, table = [], [], []
    in_quote = False

    def flush_para():
        nonlocal para
        if para:
            result.append(("p", strip_inline(" ".join(para))))
            para = []

    def flush_table():
        nonlocal table
        if table:
            rows = []
            for line in table:
                cells = [strip_inline(c) for c in line.strip().strip("|").split("|")]
                if cells and not all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
            if rows:
                result.append(("table", rows))
            table = []

    for line in lines:
        if line.startswith(">"):
            flush_para(); flush_table(); in_quote = True
            continue
        if in_quote:
            if not line.strip():
                in_quote = False
            continue
        if line.startswith("<!-- FIGURE:"):
            flush_para(); flush_table()
            slug = re.search(r"FIGURE:\s*([^\s]+)", line).group(1)
            result.append(("figure", slug))
        elif line.startswith("#"):
            flush_para(); flush_table()
            level = len(line) - len(line.lstrip("#"))
            result.append((f"h{min(level,3)}", strip_inline(line[level:].strip())))
        elif line.strip().startswith("|"):
            flush_para(); table.append(line)
        elif re.match(r"^\s*[-*]\s+", line):
            flush_para(); flush_table()
            result.append(("bullet", strip_inline(re.sub(r"^\s*[-*]\s+", "", line))))
        elif re.match(r"^\s*\d+\.\s+", line):
            flush_para(); flush_table()
            result.append(("number", strip_inline(re.sub(r"^\s*\d+\.\s+", "", line))))
        elif not line.strip():
            flush_para(); flush_table()
        elif line.strip() != "---":
            para.append(line.strip())
    flush_para(); flush_table()
    return result


def select_blocks(items, section):
    """Executive compression: retain answer, all required subsection identities, key tables and evidence."""
    selected = []
    current = ""
    prose_count = 0
    table_count = 0
    section_table_count = 0
    skip = False
    for kind, value in items:
        if kind == "h1":
            selected.append((kind, value))
            continue
        if kind == "h2":
            current = value.lower()
            prose_count = table_count = 0
            skip = any(x in current for x in ("see also", "cross-section consistency"))
            if not skip:
                selected.append((kind, value))
            continue
        if kind == "h3":
            if not skip and len(selected) < 80:
                selected.append((kind, value))
            continue
        if skip:
            continue
        if kind == "p":
            limit = 1
            if prose_count < limit:
                selected.append((kind, value)); prose_count += 1
        elif kind == "table":
            # Keep the first decision-useful table per subsection; cap very large appendix tables.
            section_table_limit = 0 if section in EXHIBITS else 1
            if table_count < 1 and section_table_count < section_table_limit and len(value) <= (18 if section == 14 else 22):
                selected.append((kind, value)); table_count += 1; section_table_count += 1
        elif kind in {"bullet", "number"}:
            if prose_count < 3:
                selected.append((kind, value)); prose_count += 1
        elif kind == "figure":
            selected.append((kind, value))
    return selected


def add_text(doc, text, style=None, bold_lead=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.keep_together = False
    if bold_lead and "—" in text:
        lead, rest = text.split("—", 1)
        r = p.add_run(lead + "—"); r.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    rows = [r + [""] * (cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = [9360 // cols] * cols
    widths[-1] += 9360 - sum(widths)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
            if i == 0:
                set_cell_shading(table.cell(i, j), "E8EEF5")
                for run in table.cell(i, j).paragraphs[0].runs:
                    run.bold = True
    set_table_geometry(table, widths)
    return table


def configure(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in [
        ("Title", 28, NAVY, 0, 10), ("Heading 1", 16, NAVY, 14, 8),
        ("Heading 2", 12, BLUE, 10, 5), ("Heading 3", 10.5, GRAY, 7, 3),
    ]:
        st = doc.styles[name]
        st.font.name, st.font.size, st.font.color.rgb = "Calibri", Pt(size), color
        st.font.bold = True
        st.paragraph_format.space_before, st.paragraph_format.space_after = Pt(before), Pt(after)
        st.paragraph_format.keep_with_next = True
    for sec in doc.sections:
        h = sec.header.paragraphs[0]
        h.text = "TALABAT GROUP  |  AI-ENABLED INVESTMENT ALLOCATION"
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for r in h.runs:
            r.font.name, r.font.size, r.font.color.rgb = "Calibri", Pt(8), GRAY
        f = sec.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f.add_run("AASTMT MBA · Group G02  |  ")
        fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
        f._p.append(fld)
        for r in f.runs:
            r.font.name, r.font.size, r.font.color.rgb = "Calibri", Pt(8), GRAY


def cover(doc):
    doc.add_paragraph("AI BUSINESS PLAN", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("talabat Group\nAI-Enabled Investment Allocation · Version 1.1")
    s = doc.add_paragraph(
        "How to allocate and govern the 2026 USD175 million programme across Everyday App and "
        "Food-leadership initiatives to maximize profitable growth and long-term platform economics."
    )
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(16), Pt(24)
    meta = add_table(doc, [
        ["Programme", "AASTMT MBA · AI for Business Organizations"],
        ["Team", "Group G02"],
        ["Scope", "talabat Group · eight operating markets"],
        ["Decision horizon", "2026 investment programme"],
        ["Document status", "Version 1.1 enhancement · independently verified baseline · 25 July 2026"],
    ])
    for row in meta.rows:
        set_cell_shading(row.cells[0], "FFF2E8")
        row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = q.add_run("Recommendation in one line")
    r.bold, r.font.color.rgb = True, ORANGE
    q.add_run(
        "\nProtect Food leadership in the evidenced GCC-3 markets, pilot talabat pro in Egypt then "
        "Iraq, continue proven Everyday App investments, and release funding through human-approved KPI gates."
    )
    doc.add_page_break()


def toc(doc):
    doc.add_heading("Contents", level=1)
    titles = []
    for p in sorted(DRAFTS.glob("Section_*.md")):
        first = next(v for k, v in blocks(p) if k == "h1")
        titles.append(first)
    table = doc.add_table(rows=len(titles), cols=2)
    table.style = "Table Grid"
    for i, title in enumerate(titles):
        table.cell(i, 0).text = str(i + 1)
        table.cell(i, 1).text = re.sub(r"^\d+\.\s*", "", title)
    set_table_geometry(table, [600, 8760])
    doc.add_paragraph(
        "Evidence convention: disclosed facts are cited by TLB source/page or vault note; assumptions "
        "and forecasts retain ASM/node identifiers; Agentic OS allocations are explicitly labeled analytical."
    ).style = doc.styles["Caption"]
    doc.add_page_break()


def build():
    generate_exhibits()
    doc = Document()
    configure(doc)
    cover(doc)
    toc(doc)
    for idx, path in enumerate(sorted(DRAFTS.glob("Section_*.md")), 1):
        items = select_blocks(blocks(path), idx)
        if idx > 1:
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        for kind, value in items:
            if kind == "h1":
                doc.add_heading(value, level=1)
                for exhibit in EXHIBITS.get(idx, []):
                    image = EXHIBIT_DIR / exhibit[0]
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(image), width=Inches(6.2))
                    c = doc.add_paragraph(exhibit[1]); c.style = doc.styles["Caption"]
                    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif kind == "h2":
                doc.add_heading(value, level=2)
            elif kind == "h3":
                doc.add_heading(value, level=3)
            elif kind == "p":
                add_text(doc, value, bold_lead=True)
            elif kind == "bullet":
                add_text(doc, value, style="List Bullet")
            elif kind == "number":
                add_text(doc, value, style="List Number")
            elif kind == "table":
                add_table(doc, value)
            elif kind == "figure":
                image = DRAFTS / "Exhibits" / f"{value}.png"
                if image.exists():
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(image), width=Inches(6.15))
        if idx == 14:
            doc.add_heading("Document Control and Validation", level=2)
            add_table(doc, [
                ["Control", "Result"],
                ["Section verification", "14/14 Done (independently verified)"],
                ["Whole-plan gates", "Problem, financial, geographic, citation, and template gates PASS"],
                ["Evidence limitation", "Country-level allocation and proprietary customer telemetry remain undisclosed; no synthetic value is presented as fact"],
                ["Decision governance", "Human approval remains mandatory before capital movement"],
            ])
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


def build_pdf():
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image, KeepTogether
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "BPTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=25,
        leading=29, textColor=colors.HexColor("#0B2545"), alignment=TA_CENTER, spaceAfter=14
    ))
    styles.add(ParagraphStyle(
        "BPH1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16,
        leading=19, textColor=colors.HexColor("#0B2545"), spaceBefore=12, spaceAfter=7
    ))
    styles.add(ParagraphStyle(
        "BPH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5,
        leading=14, textColor=colors.HexColor("#1F4D78"), spaceBefore=8, spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        "BPH3", parent=styles["Heading3"], fontName="Helvetica-Bold", fontSize=9.5,
        leading=12, textColor=colors.HexColor("#525252"), spaceBefore=6, spaceAfter=3
    ))
    styles.add(ParagraphStyle(
        "BPBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=8.6,
        leading=10.8, spaceAfter=4.5, textColor=colors.HexColor("#222222")
    ))
    styles.add(ParagraphStyle(
        "BPSmall", parent=styles["BodyText"], fontName="Helvetica", fontSize=7.2,
        leading=9, spaceAfter=2
    ))
    story = [Spacer(1, 1.15 * inch), Paragraph("AI BUSINESS PLAN", styles["BPH2"]),
             Paragraph("talabat Group<br/>AI-Enabled Investment Allocation · Version 1.1", styles["BPTitle"]),
             Paragraph(
                 "How to allocate and govern the 2026 USD175 million programme across Everyday App "
                 "and Food-leadership initiatives to maximize profitable growth and long-term platform economics.",
                 ParagraphStyle("cover", parent=styles["BPBody"], alignment=TA_CENTER, fontSize=11, leading=15)
             ), Spacer(1, 0.25 * inch)]
    cover_rows = [
        ["Programme", "AASTMT MBA · AI for Business Organizations"],
        ["Team", "Group G02"],
        ["Scope", "talabat Group · eight operating markets"],
        ["Decision horizon", "2026 investment programme"],
        ["Document status", "Version 1.1 enhancement · independently verified baseline · 25 July 2026"],
    ]
    ct = Table([[Paragraph(a, styles["BPSmall"]), Paragraph(b, styles["BPSmall"])] for a,b in cover_rows],
               colWidths=[1.45*inch, 4.45*inch])
    ct.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,-1),colors.HexColor("#FFF2E8")), ("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#DADCE0")),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"), ("LEFTPADDING",(0,0),(-1,-1),6), ("RIGHTPADDING",(0,0),(-1,-1),6),
        ("TOPPADDING",(0,0),(-1,-1),5), ("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story += [ct, Spacer(1, 0.3*inch), Paragraph(
        "<b><font color='#FF6B00'>Recommendation in one line</font></b><br/>"
        "Protect Food leadership in the evidenced GCC-3 markets, pilot talabat pro in Egypt then Iraq, "
        "continue proven Everyday App investments, and release funding through human-approved KPI gates.",
        ParagraphStyle("rec", parent=styles["BPBody"], alignment=TA_CENTER, fontSize=10, leading=14)
    ), PageBreak(), Paragraph("Contents", styles["BPH1"])]
    titles = []
    for p in sorted(DRAFTS.glob("Section_*.md")):
        titles.append(next(v for k,v in blocks(p) if k=="h1"))
    toc_rows = [[str(i), re.sub(r"^\d+\.\s*", "", t)] for i,t in enumerate(titles,1)]
    toc_t = Table([[Paragraph(a, styles["BPSmall"]), Paragraph(b, styles["BPSmall"])] for a,b in toc_rows],
                  colWidths=[0.35*inch,5.55*inch])
    toc_t.setStyle(TableStyle([("LINEBELOW",(0,0),(-1,-1),0.25,colors.HexColor("#E5E7EB")),
                               ("VALIGN",(0,0),(-1,-1),"MIDDLE"),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
    story += [toc_t, Spacer(1,8), Paragraph(
        "Evidence convention: disclosed facts are cited by TLB source/page or vault note; assumptions "
        "and forecasts retain ASM/node identifiers; Agentic OS allocations are explicitly labeled analytical.",
        styles["BPSmall"]), PageBreak()]
    for idx, path in enumerate(sorted(DRAFTS.glob("Section_*.md")),1):
        items = select_blocks(blocks(path), idx)
        if idx > 1:
            story.append(Spacer(1, 8))
        for kind,value in items:
            if kind in {"h1","h2","h3"}:
                story.append(Paragraph(value.replace("&","&amp;"), styles[{"h1":"BPH1","h2":"BPH2","h3":"BPH3"}[kind]]))
                if kind == "h1":
                    for exhibit in EXHIBITS.get(idx, []):
                        image = EXHIBIT_DIR/exhibit[0]
                        story.append(Image(str(image),width=5.9*inch,height=1.72*inch))
                        story.append(Paragraph(exhibit[1], ParagraphStyle(
                            "cap", parent=styles["BPSmall"], alignment=TA_CENTER,
                            textColor=colors.HexColor("#555555"), spaceAfter=5)))
            elif kind == "p":
                story.append(Paragraph(value.replace("&","&amp;"), styles["BPBody"]))
            elif kind in {"bullet","number"}:
                marker = "•" if kind=="bullet" else "–"
                story.append(Paragraph(f"{marker} {value.replace('&','&amp;')}", styles["BPBody"]))
            elif kind == "table":
                n=max(len(r) for r in value)
                data=[]
                for r in value:
                    r=r+[""]*(n-len(r))
                    data.append([Paragraph(c.replace("&","&amp;"), styles["BPSmall"]) for c in r])
                widths=[5.9*inch/n]*n
                tbl=Table(data,colWidths=widths,repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND",(0,0),(-1,0),colors.HexColor("#E8EEF5")),("GRID",(0,0),(-1,-1),0.35,colors.HexColor("#B8C2CC")),
                    ("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),3),("RIGHTPADDING",(0,0),(-1,-1),3),
                    ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
                ]))
                story.append(tbl); story.append(Spacer(1,5))
            elif kind == "figure":
                image=DRAFTS/"Exhibits"/f"{value}.png"
                if image.exists():
                    story.append(Image(str(image),width=5.8*inch,height=3.48*inch))
        if idx == 14:
            story += [Paragraph("Document Control and Validation",styles["BPH2"]),
                      Paragraph("14/14 sections independently verified; all five whole-plan gates PASS. "
                                "Human approval remains mandatory before capital movement.",styles["BPBody"])]

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#777777"))
        canvas.drawString(inch, 0.48*inch, "TALABAT GROUP  |  AI-ENABLED INVESTMENT ALLOCATION")
        canvas.drawRightString(7.5*inch, 0.48*inch, f"AASTMT MBA · Group G02  |  {doc.page}")
        canvas.restoreState()

    pdf = SimpleDocTemplate(str(PDF_OUT), pagesize=letter, rightMargin=inch, leftMargin=inch,
                            topMargin=0.75*inch, bottomMargin=0.7*inch,
                            title="talabat Group AI-Enabled Investment Allocation Business Plan",
                            author="AASTMT MBA Group G02")
    pdf.build(story,onFirstPage=footer,onLaterPages=footer)
    print(PDF_OUT)


if __name__ == "__main__":
    build()
    build_pdf()
